import re
from docx import Document
import requests
import json
import os

# 打开文件
# 创建一个新的Word文档
output_file = "output.docx"
doc = Document()
def fillFirm(jdata):

    headers = ["函数名", "文件名", "功能", "记述形式"]
    for header in headers:
        doc.add_table(rows=1, cols=2)
        doc.tables[-1].cell(0, 0).text = header
        doc.tables[-1].cell(0, 1).text = jdata[header]
    doc.add_table(rows=1, cols=1)
    doc.tables[-1].cell(0, 0).text = "参数"
    doc.add_table(rows=1, cols=4)
    doc.tables[-1].cell(0, 0).text = "类型"
    doc.tables[-1].cell(0, 1).text = "变量名"
    doc.tables[-1].cell(0, 2).text = "I/O"
    doc.tables[-1].cell(0, 3).text = "说明"
    for param in jdata["参数"]:
        doc.add_table(rows=1, cols=4)
        doc.tables[-1].cell(0, 0).text = param["类型"]
        doc.tables[-1].cell(0, 1).text = param["变量名"]
        doc.tables[-1].cell(0, 2).text = param["I/O"]
        doc.tables[-1].cell(0, 3).text = param["说明"]
    retNotNUll=False
    for ret in jdata["返回值"]:
        retNotNUll=True
        doc.add_table(rows=1, cols=4)
        doc.tables[-1].cell(0, 0).text = "返回值"
        doc.tables[-1].cell(0, 1).text = "型"
        doc.tables[-1].cell(0, 2).text = ret["类型"]
        doc.tables[-1].cell(0, 3).text = "说明"
        for value in ret["值"]:
            doc.add_table(rows=1, cols=4)
            doc.tables[-1].cell(0, 1).text = "值"
            doc.tables[-1].cell(0, 2).text = value["值"]
            doc.tables[-1].cell(0, 3).text = value["说明"]
    if retNotNUll==False:
        doc.add_table(rows=1, cols=4)
        doc.tables[-1].cell(0, 0).text = "返回值"
        doc.tables[-1].cell(0, 1).text = "型"
        # doc.tables[-1].cell(0, 2).text = ret["类型"]
        doc.tables[-1].cell(0, 3).text = "说明"
    doc.add_table(rows=1, cols=1)
    doc.tables[-1].cell(0, 0).text = "详细说明"
    doc.add_table(rows=1, cols=1)
    doc.tables[-1].cell(0, 0).text = jdata["详细说明"]
    doc.add_paragraph("\n\n")


class deepseekAPI:
    def __init__(self):
        self.url = "https://api.deepseek.com/chat/completions"

        # 请求头
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": "Bearer <DeepSeek API>"  # 替换为你的 DeepSeek API 密钥
        }


    def requestAPI(self,prompt):
        # 请求数据
        self.data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            "stream": False
        }
        response = requests.post(self.url, headers=self.headers, json=self.data)
        return response.status_code,response.json()["choices"][0]["message"]["content"]

if __name__ == '__main__':
    deepseek = deepseekAPI()
    for root, dirs, files in os.walk('./'):
        for file in files:
            # 检查文件扩展名
            if file.endswith('.vue'):
                with open(os.path.join(root, file), 'r', encoding='utf-8') as f:
                    print(f"正在处理文件: {file}")
                    content = f.read()

                    prompt = content+'''上面是我的代码，下面是一个json，请你根据示范为我的代码生成一个json，不要包含多余的描述
                [
                {
                    "函数名": "requestSendImg",
                    "文件名": "HandwritingRecognition.vue",
                    "功能": "请求发送图片",
                    "记述形式": "bool requestSendImg(quint32 receiver, QString imgAbsolutePath);",
                    "参数": [
                        {
                            "类型": "quint32",
                            "变量名": "receiver",
                            "I/O": "I",
                            "说明": "接收者"
                        },
                        {
                            "类型": "QString",
                            "变量名": "imgAbsolutePath",
                            "I/O": "I",
                            "说明": "图片路径"
                        }
                    ],
                    "返回值": [{
                        "类型": "bool",
                        "值": [
                            {
                                "值": "true",
                                "说明": "请求成功"
                            },
                            {
                                "值": "false",
                                "说明": "请求失败"
                            }
                        ]
                    },
                    ],
                    "详细说明": "请求发送图片，将图片缓存到本地。"
                },
                {
                    "函数名": "handleButtonClick_submit",
                    "文件名": "HandwritingRecognition.vue",
                    "功能": "提交手写数字识别请求",
                    "记述形式": "handleButtonClick_submit()",
                    "参数": [],
                    "返回值": [],
                    "详细说明": "获取CanvasDrawingBoard组件中的画布图像，将其转换为Base64编码并发送到服务器进行手写数字识别。根据识别模式（单数字或连续数字）处理返回结果，并在页面上显示识别结果或 更新画布内容。"
                },
                ]'''
                    status_code, response = deepseek.requestAPI(prompt)
                    # 使用正则表达式查找```和```之间的内容
                    # matches = re.findall(r'```(.*?)```', response, re.DOTALL)
                    print(response[7:-3])
                    # 将字符串转换为JSON对象
                    json_obj = json.loads(response[7:-3])
                    for item in json_obj:
                        fillFirm(item)
        
    # 保存Word文档
    doc.save(output_file)
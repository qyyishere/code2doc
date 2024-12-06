from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from zhconv import convert
from docx.shared import Pt
graylist = ["函数名","文件名","功能","记述形式","参数","类型","变量名","I/O","说明","返回值","型","说明","详细说明"]
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

def filter_paragraphs(doc):
    # # 遍历文档中的所有表格
    # for table in doc.tables:
    #     # 遍历表格中的所有行
    #     for row in table.rows:
    #         # 遍历行中的所有单元格
    #         for cell in row.cells:
    #             # 遍历单元格中的所有段落
    #             for paragraph in cell.paragraphs:
    #                 # 检查段落文本是否以“4.1”开头
    #                 if not paragraph.text.startswith("4.1"):
    #                     # 删除段落
    #                     paragraph._element.getparent().remove(paragraph._element)

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 检查段落文本是否以“4.1”开头
        if not paragraph.text.startswith("4.1"):
            # 删除段落
            print(paragraph.text)
            delete_paragraph(paragraph)# doc.paragraphs.remove(paragraph)


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_border_to_tables(doc):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 为表格添加所有框线
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell,
            top={"sz": 5, "val": "single", "color": "#000000", },
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            left={"sz": 5, "color": "#000000","val": "single"},
            right={"sz": 5, "color": "#000000","val": "single"},)

def convert_to_simplified(doc):

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 获取段落的文本
        text = paragraph.text
        # 更新段落的文本
        paragraph.text = convert(text,'zh-cn')
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的所有行
        for row in table.rows:
            # 遍历行中的所有单元格
            for cell in row.cells:
                # 获取单元格中的文本
                text = cell.text
                # 更新单元格的文本
                cell.text = convert(text,'zh-cn')

def set_cell_background(doc):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的所有行
        for row in table.rows:
            # 遍历行中的所有单元格
            for cell in row.cells:
                # 获取单元格中的文本
                cell_text = cell.text.strip()
                # 如果文本是“屬性”或“函數”，则设置背景颜色为灰色
                if cell_text in graylist:
                    # 获取单元格的XML元素  
                    tc = cell._tc  
                    tcPr = tc.get_or_add_tcPr()  
                    
                    # 设置单元格背景颜色为  
                    shd = OxmlElement('w:shd')  
                    shd.set(qn('w:val'), 'pct100')  # 设置填充类型为100%填充  
                    shd.set(qn('w:fill'), '#CCCCCC')  # 设置填充颜色为（RGB）  
                    
                    # 移除现有的shd元素（如果有的话）  
                    existing_shds = tcPr.xpath('.//w:shd')  
                    for existing_shd in existing_shds:  
                        tcPr.remove(existing_shd)  
                    
                    # 添加新的背景色元素到单元格属性中  
                    tcPr.append(shd) 


def set_font(doc):
    # 设置中文字体和字号
    chinese_font = '宋体'
    chinese_size = Pt(10)
    # 设置英文字体和字号
    english_font = 'Times New Roman'
    english_size = Pt(10)
    
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 遍历段落中的所有运行
        for run in paragraph.runs:
            # 获取运行中的文本
            text = run.text
            # 如果文本包含中文字符，则设置字体和字号
            if any('\u4e00' <= char <= '\u9fff' for char in text):
                run.font.name = chinese_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                run.font.size = chinese_size
            # 如果文本包含英文字符，则设置字体和字号
            elif any('a' <= char <= 'z' or 'A' <= char <= 'Z' for char in text):
                run.font.name = english_font
                run.font.size = english_size
    
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的所有行
        for row in table.rows:
            # 遍历行中的所有单元格
            for cell in row.cells:
                # 遍历单元格中的所有运行
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # 获取运行中的文本
                        text = run.text
                        # 如果文本包含中文字符，则设置字体和字号
                        if any('\u4e00' <= char <= '\u9fff' for char in text):
                            run.font.name = chinese_font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                            run.font.size = chinese_size
                        # 如果文本包含英文字符，则设置字体和字号
                        elif any('a' <= char <= 'z' or 'A' <= char <= 'Z' for char in text):
                            run.font.name = english_font
                            run.font.size = english_size
    

if __name__ == "__main__":
    str_doc = f"C:/Users/qyyis/OneDrive/Desktop/workspace/output.docx"
    # 打开Word文档
    doc = Document(str_doc)
    set_cell_background(doc)
    add_border_to_tables(doc)
    #rgb(204,204,204)
    #convert_to_simplified(doc)
    set_font(doc)
    # 保存修改后的文档
    doc.save(f"output.docx")
    
